from docx import Document
import os

def generate_word(single_report=None):
    document = Document()

    updates = [single_report] if single_report else Update.query.order_by(Update.week.desc()).all()

    for update in updates:
        document.add_heading(f"{update.name} ({update.position}) - {update.week}", level=1)
        document.add_paragraph(f"프로젝트 정보:\n{update.project_summary}")
        document.add_paragraph(f"Key Milestone:\n{update.milestones}")
        document.add_paragraph(f"진행상황:\n{update.progress}")
        document.add_paragraph(f"특이사항:\n{update.project_issues}")
        document.add_paragraph(f"영업지원:\n{update.sales_support}")
        document.add_paragraph(f"기타 특이사항:\n{update.other_notes}")
        document.add_paragraph(f"출장:\n{update.business_trip}")
        document.add_paragraph(f"외근:\n{update.external_work}")
        document.add_paragraph(f"휴가:\n{update.vacation}")
        document.add_paragraph(f"휴일근무:\n{update.weekend_work}")
        document.add_paragraph("\n")

    output_path = os.path.join("instance", "report_export.docx")
    document.save(output_path)
    return output_path
