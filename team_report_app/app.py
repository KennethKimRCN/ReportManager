import os
from flask import (
    Flask, render_template, redirect, url_for,
    request, session, flash, send_file
)
from models import db, User, Project, Update
from functools import wraps
from io import BytesIO
from docx import Document

basedir = os.path.abspath(os.path.dirname(__file__))
db_path = os.path.join(basedir, 'instance', 'updates.db')

app = Flask(__name__)
app.secret_key = 'change_this_secret_to_a_random_value'
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

# --- Login required decorator ---
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'employee_id' not in session:
            flash("로그인이 필요합니다.")
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# --- LOGIN ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        employee_id = request.form.get('employee_id')
        user = User.query.filter_by(employee_id=employee_id).first()
        if user:
            session['employee_id'] = user.employee_id
            session['name'] = user.name
            session['position'] = user.position
            return redirect(url_for('dashboard'))
        else:
            error = '존재하지 않는 사번입니다.'
    return render_template('login.html', error=error)

# --- LOGOUT ---
@app.route('/logout')
@login_required
def logout():
    session.clear()
    flash('로그아웃 되었습니다.')
    return redirect(url_for('login'))

# --- DASHBOARD ---
@app.route('/')
@login_required
def dashboard():
    reports = Update.query.order_by(Update.week.desc()).all()
    return render_template('dashboard.html', reports=reports)

# --- CREATE PROJECT (popup modal) ---
@app.route('/projects/new', methods=['POST'])
@login_required
def create_project():
    solution_name = request.form['solution_name']
    location = request.form['location']
    company = request.form['company']
    project_name = request.form['project_name']
    code = request.form['code']

    user = User.query.filter_by(employee_id=session['employee_id']).first()
    if not user:
        flash("사용자를 찾을 수 없습니다.")
        return redirect(url_for('dashboard'))

    project = Project(
        solution_name=solution_name,
        location=location,
        company=company,
        project_name=project_name,
        code=code
    )
    project.assignees.append(user)

    db.session.add(project)
    db.session.commit()
    flash("프로젝트가 등록되었습니다.")
    return redirect(url_for('dashboard'))

# --- SUBMIT REPORT ---
@app.route('/submit', methods=['GET', 'POST'])
@login_required
def submit():
    if request.method == 'POST':
        employee_id = session['employee_id']
        week = request.form['week']
        project_ids = request.form.getlist('project_ids')

        user = User.query.filter_by(employee_id=employee_id).first()
        if not user:
            flash("사용자를 찾을 수 없습니다.")
            return redirect(url_for('submit'))

        for pid in project_ids:
            project = Project.query.get(pid)
            if not project:
                continue

            update = Update(
                user_id=user.id,
                project_id=project.id,
                week=week,
                progress=request.form.get(f'project_{pid}_progress'),
                project_issues=request.form.get(f'project_{pid}_issues'),
                sales_support=request.form.get(f'project_{pid}_sales_support'),
                other_notes=request.form.get(f'project_{pid}_other_notes'),
                business_trip=request.form.get(f'project_{pid}_business_trip'),
                external_work=request.form.get(f'project_{pid}_external_work'),
                vacation=request.form.get(f'project_{pid}_vacation'),
                weekend_work=request.form.get(f'project_{pid}_weekend_work')
            )
            db.session.add(update)

        db.session.commit()
        flash("보고서가 저장되었습니다.")
        return redirect(url_for('dashboard'))

    projects = Project.query.order_by(Project.solution_name).all()
    return render_template('submit.html', projects=projects)

# --- VIEW REPORT ---
@app.route('/report/<int:update_id>')
@login_required
def view_report(update_id):
    report = Update.query.get_or_404(update_id)
    return render_template('view_report.html', report=report)

# --- EDIT REPORT ---
@app.route('/edit/<int:update_id>', methods=['GET', 'POST'])
@login_required
def edit_report(update_id):
    update = Update.query.get_or_404(update_id)

    user = User.query.filter_by(employee_id=session['employee_id']).first()
    if update.user_id != user.id:
        flash("권한이 없습니다.")
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        update.week = request.form['week']
        update.progress = request.form.get('progress')
        update.project_issues = request.form.get('project_issues')
        update.sales_support = request.form.get('sales_support')
        update.other_notes = request.form.get('other_notes')
        update.business_trip = request.form.get('business_trip')
        update.external_work = request.form.get('external_work')
        update.vacation = request.form.get('vacation')
        update.weekend_work = request.form.get('weekend_work')

        db.session.commit()
        flash("보고서가 수정되었습니다.")
        return redirect(url_for('view_report', update_id=update.id))

    weeks = [f"Y25W{str(i).zfill(2)}" for i in range(1, 53)]
    return render_template('edit_report.html', update=update, weeks=weeks)

# --- DELETE REPORT ---
@app.route('/delete/<int:update_id>', methods=['POST'])
@login_required
def delete_report(update_id):
    update = Update.query.get_or_404(update_id)
    user = User.query.filter_by(employee_id=session['employee_id']).first()
    if update.user_id != user.id:
        flash("권한이 없습니다.")
        return redirect(url_for('dashboard'))

    db.session.delete(update)
    db.session.commit()
    flash("보고서가 삭제되었습니다.")
    return redirect(url_for('dashboard'))

# --- EXPORT TO WORD ---
@app.route('/export/<int:update_id>')
@login_required
def export_report(update_id):
    update = Update.query.get_or_404(update_id)

    doc = Document()
    doc.add_heading(f"[주간보고] {update.user.name} {update.user.position} – {update.week}", level=1)

    def add_section(title, content):
        doc.add_paragraph(f"▶ {title}")
        if content and content.strip():
            doc.add_paragraph(content)
        else:
            doc.add_paragraph("내용 없음")

    add_section("프로젝트 특이사항", update.project_issues)
    add_section("Key Milestone", "")
    add_section("진행상황", update.progress)
    add_section("특이사항", update.project_issues)
    add_section("영업지원 사항", update.sales_support)
    add_section("그 외 특이사항", update.other_notes)

    doc.add_paragraph("▶ 개인 일정")
    doc.add_paragraph(f"출장: {update.business_trip or ''}")
    doc.add_paragraph(f"외근: {update.external_work or ''}")
    doc.add_paragraph(f"휴가: {update.vacation or ''}")
    doc.add_paragraph(f"휴일근무: {update.weekend_work or ''}")

    f = BytesIO()
    doc.save(f)
    f.seek(0)

    filename = f"{update.week}_{update.user.name}.docx"
    return send_file(
        f,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/admin/data')
@login_required
def admin_data():
    users = User.query.all()
    projects = Project.query.all()
    updates = Update.query.all()
    return render_template('admin_data.html', users=users, projects=projects, updates=updates)

if __name__ == '__main__':
    app.run(debug=True)
